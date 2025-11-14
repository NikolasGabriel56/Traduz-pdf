from pdf2docx import Converter
import docx
from docx import Document
from deep_translator import GoogleTranslator
from docx2pdf import convert
from docx.shared import Pt
from docx.oxml.ns import qn
import os
from tkinter import Tk, filedialog, messagebox, StringVar, Label, Button, Entry


# ============================================================
#  ETAPA 1 — Converter PDF → DOCX preservando layout original
# ============================================================
def pdf_para_docx(pdf_path):
    docx_path = pdf_path.replace(".pdf", ".docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    return docx_path


# ============================================================
#  ETAPA 2 — Traduzir texto mantendo layout de TABELAS IGUAL
# ============================================================
def traduzir_docx(docx_path, idioma_destino):
    tradutor = GoogleTranslator(source="auto", target=idioma_destino)
    doc = Document(docx_path)

    # Margens bem discretas (não alteram visual)
    CM = 360000
    for section in doc.sections:
        section.top_margin = int(0.5 * CM)
        section.bottom_margin = int(0.5 * CM)
        section.left_margin = int(0.8 * CM)
        section.right_margin = int(0.8 * CM)

    # =============================
    # TRADUÇÃO DE PARÁGRAFOS (fora das tabelas)
    # Preservando runs
    # =============================
    for p in doc.paragraphs:
        if not p.text.strip():
            continue

        texto_original = p.text
        try:
            texto_traduzido = tradutor.translate(texto_original)
        except:
            continue

        # salvar runs originais
        estilos_runs = []
        for run in p.runs:
            estilos_runs.append({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size
            })

        # limpar parágrafo
        while p.runs:
            p.runs[0].clear()
            p._element.remove(p.runs[0]._element)

        # recriar novo run com formatação original (melhor possível)
        run_novo = p.add_run(texto_traduzido)

        if estilos_runs:
            run_novo.bold = estilos_runs[0]["bold"]
            run_novo.italic = estilos_runs[0]["italic"]
            run_novo.underline = estilos_runs[0]["underline"]
            if estilos_runs[0]["font_name"]:
                run_novo.font.name = estilos_runs[0]["font_name"]
            if estilos_runs[0]["font_size"]:
                run_novo.font.size = estilos_runs[0]["font_size"]

        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0

    # =============================
    # TABELAS — tradução com preservação total de formatação
    # =============================
    for tabela in doc.tables:
        for row in tabela.rows:

            # Altura automática da linha
            trPr = row._tr.get_or_add_trPr()
            autoHeight = docx.oxml.parse_xml(
                r"<w:trHeight w:hRule='auto' xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'/>"
            )
            trPr.append(autoHeight)

            for celula in row.cells:

                texto_original = celula.text.strip()
                if not texto_original:
                    continue

                try:
                    texto_traduzido = tradutor.translate(texto_original)
                except:
                    continue

                for para in celula.paragraphs:

                    # salva os estilos de todos os runs
                    estilos_runs = []
                    for run in para.runs:
                        estilos_runs.append({
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_name": run.font.name,
                            "font_size": run.font.size
                        })

                    # limpar o conteúdo mantendo formatação estrutural
                    para.clear()

                    # recriar run(s)
                    run_novo = para.add_run(texto_traduzido)

                    # aplicar estilo original do primeiro run da célula
                    if estilos_runs:
                        run_novo.bold = estilos_runs[0]["bold"]
                        run_novo.italic = estilos_runs[0]["italic"]
                        run_novo.underline = estilos_runs[0]["underline"]

                        if estilos_runs[0]["font_name"]:
                            run_novo.font.name = estilos_runs[0]["font_name"]

                        if estilos_runs[0]["font_size"]:
                            run_novo.font.size = estilos_runs[0]["font_size"]

                    para.paragraph_format.line_spacing = 1
                    para.paragraph_format.space_after = 0

    novo_docx = docx_path.replace(".docx", f"_{idioma_destino}.docx")
    doc.save(novo_docx)
    return novo_docx

# ============================================================
#  ETAPA 3 — Converter DOCX traduzido → PDF
# ============================================================
def docx_para_pdf(docx_path):
    pasta = os.path.dirname(docx_path)
    convert(docx_path, pasta)
    return docx_path.replace(".docx", ".pdf")


# ============================================================
#  FLUXO PRINCIPAL
# ============================================================
def traduzir_pdf_layout_total(pdf_path, idioma_destino):
    docx_temp = pdf_para_docx(pdf_path)
    traduzido_docx = traduzir_docx(docx_temp, idioma_destino)
    saida_pdf = docx_para_pdf(traduzido_docx)
    print(f"✅ PDF final salvo em: {saida_pdf}")
    return saida_pdf


# ============================================================
#  INTERFACE TKINTER
# ============================================================
def selecionar_arquivo():
    arquivo = filedialog.askopenfilename(
        title="Selecione um PDF",
        filetypes=[("Arquivos PDF", "*.pdf")]
    )
    caminho_var.set(arquivo)


def traduzir():
    caminho = caminho_var.get().strip()
    idioma = idioma_var.get().strip()

    if not caminho or not idioma:
        messagebox.showerror("Erro", "Selecione o PDF e o idioma de destino.")
        return

    try:
        saida = traduzir_pdf_layout_total(caminho, idioma)
        messagebox.showinfo("Concluído", f"Tradução concluída!\nPDF salvo em:\n{saida}")
    except Exception as e:
        messagebox.showerror("Erro", f"Falha durante a tradução:\n{e}")


# ============================================================
#  JANELA
# ============================================================
root = Tk()
root.title("Tradutor de PDFs — SERTA ⚙️ (Layout Preservado)")
root.geometry("560x300")
root.resizable(False, False)

caminho_var = StringVar()
idioma_var = StringVar()

Label(root, text="Selecione o PDF:", font=("Arial", 11)).pack(pady=10)
Entry(root, textvariable=caminho_var, width=65).pack()
Button(root, text="Escolher PDF", command=selecionar_arquivo).pack(pady=5)

Label(root, text="Idioma destino (ex: en, es, fr, it):", font=("Arial", 11)).pack(pady=10)
Entry(root, textvariable=idioma_var, width=15, font=("Arial", 12)).pack()

Button(root, text="Traduzir", command=traduzir,
       bg="#4CAF50", fg="white", font=("Arial", 12, "bold")).pack(pady=20)

root.mainloop()
